# poetic_waveform_lab_v2_11_integra_font_fixed.py

# -*- coding: utf-8 -*-
# poetic_waveform_lab_v2_11_integra_font_fixed.py

import os
import re
import io
import time
import hashlib
from datetime import datetime

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib import font_manager, rcParams
from matplotlib.font_manager import FontProperties

import streamlit as st
import docx  # for reading .docx
from docx import Document  # for writing .docx
from docx.shared import Inches

# ============================ OpenAI client (v1 / legacy) =====================
try:
    from openai import OpenAI
    _OPENAI_CLIENT_MODE = "v1"
except Exception:  # 旧SDK
    import openai as _openai_legacy
    _OPENAI_CLIENT_MODE = "legacy"

def _get_openai_client():
    """Return an OpenAI client that works for both v1 and legacy SDKs."""
    # Prefer Streamlit secrets, then env var
    api_key = None
    try:
        api_key = st.secrets["openai"]["api_key"]
    except Exception:
        api_key = os.getenv("OPENAI_API_KEY", "")

    if _OPENAI_CLIENT_MODE == "v1":
        client = OpenAI(api_key=api_key) if api_key else OpenAI()
        return client, "v1"
    else:
        _openai_legacy.api_key = api_key
        return _openai_legacy, "legacy"

# ============================ Embedder ========================================
@st.cache_resource(show_spinner=False)
def load_embedder():
    """Load sentence-transformer; return None on failure but keep app alive."""
    try:
        from sentence_transformers import SentenceTransformer
        return SentenceTransformer("all-MiniLM-L6-v2")
    except Exception as e:
        st.warning(f"埋め込みモデルを読み込めませんでした。Divergenceはゼロ系列で継続します。詳細: {e}")
        return None

# ============================ Font setup ======================================
def setup_japanese_font(font_path: str = "") -> str:
    """
    日本語フォントをMatplotlibに登録し、PDFで埋め込まれるよう設定。
    戻り値: 実際にセットされたフォントファミリ名（空なら未適用）。
    """
    rcParams["pdf.fonttype"] = 42
    rcParams["ps.fonttype"]  = 42
    rcParams["pdf.use14corefonts"] = False
    rcParams["axes.unicode_minus"] = False
    rcParams["text.usetex"] = False

    used_name = ""

    def _use(fp: FontProperties):
        nonlocal used_name
        used_name = fp.get_name()
        # family名と各分類に同名を入れて“強制適用”
        rcParams["font.family"] = used_name
        rcParams["font.sans-serif"] = [used_name]
        rcParams["font.serif"] = [used_name]
        return used_name

    # 1) 明示パス優先
    if font_path and os.path.exists(font_path):
        try:
            font_manager.fontManager.addfont(font_path)
            return _use(FontProperties(fname=font_path))
        except Exception:
            pass

    # 2) リポ同梱探索
    for candidate in [
        "fonts/NotoSansJP-Regular.ttf",
        "fonts/NotoSerifJP-Regular.ttf",
        "fonts/NotoSansJP-Bold.ttf",
        "fonts/NotoSerifJP-Bold.ttf",
    ]:
        if os.path.exists(candidate):
            try:
                font_manager.fontManager.addfont(candidate)
                return _use(FontProperties(fname=candidate))
            except Exception:
                continue

    # 3) OSインストール名から
    installed = {f.name for f in font_manager.fontManager.ttflist}
    for name in [
        "Noto Sans JP", "Noto Serif JP",
        "IPAexGothic", "IPAexMincho", "Yu Gothic", "Meiryo",
        "Hiragino Sans", "Hiragino Mincho ProN"
    ]:
        if name in installed:
            return _use(FontProperties(family=name))

    # 見つからなければ欧文のまま（PDFで豆腐化する可能性は残る）
    return used_name

def _apply_font_to_fig(fig, family_name: str):
    """既に作成済みの Figure 内の全テキストにフォントを適用"""
    fp = FontProperties(family=family_name)
    for ax in fig.get_axes():
        # 軸ラベル・タイトル・凡例・目盛り
        if ax.title:
            ax.title.set_fontproperties(fp)
        if ax.xaxis.label:
            ax.xaxis.label.set_fontproperties(fp)
        if ax.yaxis.label:
            ax.yaxis.label.set_fontproperties(fp)
        for t in ax.get_xticklabels() + ax.get_yticklabels():
            t.set_fontproperties(fp)
        leg = ax.get_legend()
        if leg:
            for t in leg.get_texts():
                t.set_fontproperties(fp)
        # 手動注釈
        for txt in ax.texts:
            txt.set_fontproperties(fp)

# ============================ Utilities =======================================
def extract_lines_from_docx(file):
    """Extract non-empty lines from a .docx treating each paragraph as a line."""
    try:
        doc = docx.Document(file)
    except Exception as e:
        st.error(f"docx読み込みでエラー: {e}")
        return []
    lines = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            lines.append(text)
    return lines

def _safe_minmax_scale(arr: np.ndarray) -> np.ndarray:
    """Scale to [0,1]; return zeros if constant or empty."""
    if arr.size == 0:
        return arr
    mn, mx = float(np.min(arr)), float(np.max(arr))
    if mx - mn == 0:
        return np.zeros_like(arr, dtype=float)
    return (arr - mn) / (mx - mn)

def _clean_list_lines(raw: str, limit: int = 10):
    """
    Normalize LLM list output -> list[str]:
    - split by lines or commas
    - strip bullets / numbers / dashes
    - drop empties, dedupe, limit
    """
    if not raw:
        return []
    parts = []
    for line in re.split(r"[\n,]+", raw):
        s = re.sub(r"^\s*(?:-|\*|\d+[\.\)]|\u2022)\s*", "", line.strip())
        s = s.strip("・:：;、。 \t")
        if s:
            parts.append(s)
    # Deduplicate preserving order
    seen, out = set(), []
    for p in parts:
        if p not in seen:
            seen.add(p)
            out.append(p)
        if len(out) >= limit:
            break
    return out

def _hash_for_cache(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()

# ============================ Scoring =========================================
def _cos_divergences(embeddings: np.ndarray, window: int) -> np.ndarray:
    """Compute 1 - cosine(line_i, mean(prev window)) without normalization."""
    n = embeddings.shape[0]
    scores = np.zeros(n, dtype=float)
    for i in range(n):
        if i < window:
            scores[i] = 0.0
            continue
        ctx = embeddings[i-window:i].mean(axis=0)
        a = embeddings[i]
        denom = (np.linalg.norm(a) * np.linalg.norm(ctx)) or 1e-9
        cos = float(np.dot(a, ctx) / denom)
        scores[i] = 1.0 - cos
    return scores

def compute_divergence_scores(lines, model, window: int = 3, normalize: bool = True):
    """Compute divergence; optionally min-max normalize."""
    if not lines:
        return np.array([], dtype=float)
    if model is None:
        return np.zeros(len(lines), dtype=float)
    try:
        embeddings = model.encode(lines)
    except Exception as e:
        st.warning(f"埋め込み計算に失敗: {e}（Divergenceはゼロで継続）")
        return np.zeros(len(lines), dtype=float)
    raw = _cos_divergences(np.asarray(embeddings), max(1, int(window)))
    return _safe_minmax_scale(raw) if normalize else raw

def compute_difficulty_scores(lines, difficult_words, decay=0.5):
    """Afterglow accumulation of GPT-extracted hard-word hits with safe scaling."""
    n = len(lines)
    if n == 0:
        return np.array([], dtype=float)
    difficult_words = [w for w in (difficult_words or []) if w]
    prev = 0.0
    scores = []
    for line in lines:
        base = 0
        if difficult_words:
            base = sum((w in line) for w in difficult_words)
        prev = decay * prev + float(base)
        scores.append(prev)
    return _safe_minmax_scale(np.array(scores, dtype=float))

def compute_resonance_scores_from_terms(lines, resonance_terms, decay=0.5):
    """Resonance scoring with afterglow (余韻)."""
    n = len(lines)
    if n == 0:
        return np.array([], dtype=float)
    terms = set(resonance_terms or [])
    prev = 0.0
    scores = []
    for line in lines:
        base = 0
        if terms:
            base = sum((w in line) for w in terms)
        prev = decay * prev + float(base)
        scores.append(prev)
    return _safe_minmax_scale(np.array(scores, dtype=float))

# ============================ GPT helper ======================================
@st.cache_data(show_spinner=False)
def _cached_gpt_words(cache_key: str, system_prompt: str, user_prompt: str, model_name: str, mode: str, retries: int = 3):
    """Cache wrapper for GPT extraction with simple retry/backoff."""
    client, _ = _get_openai_client()

    def _call():
        if mode == "v1":
            resp = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.3,
            )
            return resp.choices[0].message.content
        else:
            resp = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": user_prompt}],
                temperature=0.3,
            )
            msg = resp.choices[0].message
            return msg.get("content") if isinstance(msg, dict) else getattr(msg, "content", "")

    last_err = None
    for i in range(retries):
        try:
            return _call()
        except Exception as e:
            last_err = e
            time.sleep(0.7 * (2 ** i))
    raise RuntimeError(f"GPT抽出に失敗: {last_err}")

def gpt_extract_words(prompt: str, model_name: str = "gpt-4o-mini") -> list:
    client, mode = _get_openai_client()
    sys_prompt = "出力は箇条書きや番号を付けず、語のみを改行区切りで返してください。句読点や説明は不要です。最大10語。"
    cache_key = _hash_for_cache(model_name + "||" + prompt)
    try:
        raw = _cached_gpt_words(cache_key, sys_prompt, prompt, model_name, mode)
    except Exception as e:
        st.info(f"GPT抽出をスキップ（{e}）: 空集合で継続します。")
        return []
    return _clean_list_lines(raw, limit=10)

# ============================ Report builders =================================
def build_docx_report(lines, df, difficult_words, resonance_words, fig_wave, fig_quad, title="Poetic Waveform Report"):
    """スコア表・波形・四象限を1つのDOCXにまとめて返す（BytesIO）。"""
    buf = io.BytesIO()
    doc = Document()
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {ts}")

    # パラメータ概要
    doc.add_heading("Parameters", level=2)
    doc.add_paragraph(f"Lines: {len(lines)}")

    # 語群
    doc.add_heading("Difficult terms (GPT)", level=3)
    doc.add_paragraph(", ".join(difficult_words or []))
    doc.add_heading("Resonant terms (GPT)", level=3)
    doc.add_paragraph(", ".join(resonance_words or []))

    # スコア表
    doc.add_heading("Scores", level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            cells[j].text = str(row[col])

    # 画像（fig を一旦PNGに落として挿入）
    def _add_fig(pfig, caption):
        buf_img = io.BytesIO()
        pfig.savefig(buf_img, format="png", bbox_inches="tight", dpi=180)
        buf_img.seek(0)
        doc.add_heading(caption, level=2)
        doc.add_picture(buf_img, width=Inches(6.2))
        buf_img.close()

    _add_fig(fig_wave, "Waveforms (Divergence / Difficulty / Resonance)")
    _add_fig(fig_quad, "Quadrant (raw) with Centroid")

    doc.save(buf)
    buf.seek(0)
    return buf
 

    # 依存を増やさず Matplotlib だけでPDFを構成（表は分割して簡易表示）。
    # ページ1: 表紙／概要、ページ2: 波形、ページ3: 四象限、ページ4+: スコア表
 

def build_pdf_report(lines, df, fig_wave, fig_quad, title="Poetic Waveform Report", font_family: str = ""):
    pdf_buf = io.BytesIO()
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")   # ←追加
    with PdfPages(pdf_buf) as pdf:

        # 1) 使うフォント名を決定（genericは回避）
        family = font_family or rcParams.get("font.family", "")
        if isinstance(family, (list, tuple)):
            family = family[0] if family else ""
        if not family or family in ("sans-serif", "serif", "monospace"):
            # 最終フォールバック：Matplotlib 同梱の DejaVu で落とさず出す（日本語は□になる可能性）
            family = "DejaVu Sans"
        fp = FontProperties(family=family)

        # ……以降は現行どおり……
        # 表紙の ax0.text(..., fontproperties=fp)
        # _apply_font_to_fig(fig_wave, family)
        # _apply_font_to_fig(fig_quad, family)
        # テーブルの cell.get_text().set_fontproperties(fp)

       
        # 表紙
        fig0, ax0 = plt.subplots(figsize=(8.27, 11.69))  # A4縦相当
        ax0.axis("off")
        text = f"""{title}
Generated: {ts}




Lines: {len(lines)}
Columns: {', '.join(df.columns)}
"""
        ax0.text(0.05, 0.95, text, va="top", fontsize=14, fontproperties=fp)
        pdf.savefig(fig0, bbox_inches="tight")
        plt.close(fig0)

        # 既存図にもフォントを強制（凡例/目盛り/注釈まで）
        _apply_font_to_fig(fig_wave, family)
        _apply_font_to_fig(fig_quad, family)

        # 波形
        pdf.savefig(fig_wave, bbox_inches="tight")

        # 四象限
        pdf.savefig(fig_quad, bbox_inches="tight")

        # スコア表（分割して複数ページ）
        rows_per_page = 35
        for start in range(0, len(df), rows_per_page):
            chunk = df.iloc[start:start + rows_per_page]
            figt, axt = plt.subplots(figsize=(8.27, 11.69))
            axt.axis("off")
            axt.set_title(f"Scores (rows {start+1}–{start+len(chunk)})", pad=12, fontproperties=fp)

            tbl = axt.table(cellText=chunk.values,
                            colLabels=list(chunk.columns),
                            loc="center")
            tbl.auto_set_font_size(False)
            tbl.set_fontsize(8)
            tbl.scale(1.0, 1.2)
            # セルにフォントを厳密適用
            for cell in tbl.get_celld().values():
                cell.get_text().set_fontproperties(fp)

            pdf.savefig(figt, bbox_inches="tight")
            plt.close(figt)

    pdf_buf.seek(0)
    return pdf_buf

# ============================ UI ==============================================
st.set_page_config(page_title="詩的波形分析室 v2.11 — 統合・堅牢版", layout="wide")
st.title("詩的波形分析室 v2.11 — Resonance余韻 + Divergence切替 + 四象限（raw/重心）")

# 起動直後（set_page_config のすぐ下）あたりに
FONT_PATH = "fonts/NotoSansJP-Regular.ttf"  # 置いた実ファイル名に合わせて
if not os.path.exists(FONT_PATH):
    st.error(f"指定フォントが見つかりません: {FONT_PATH}  （fonts/ 配下に ttf を置いてください）")
used_family = setup_japanese_font(FONT_PATH)
st.caption(f"PDF用フォント: {used_family or '未セット（フォールバック）'}")

# ---- 入力モード --------------------------------------------------------------
mode = st.radio("入力モードを選択", ["DOCXアップロード", "テキスト貼り付け"], horizontal=True)

lines = []
if mode == "DOCXアップロード":
    uploaded_file = st.file_uploader("詩の .docx ファイルをアップロードしてください", type=["docx"])
    if uploaded_file:
        lines = extract_lines_from_docx(uploaded_file)
elif mode == "テキスト貼り付け":
    raw_text = st.text_area("詩本文（行ごとに改行）", height=220, placeholder="ここに本文を貼り付け…")
    if raw_text and raw_text.strip():
        lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]

if not lines:
    st.info("入力が未指定です。DOCX をアップロードするか、本文を貼り付けてください。")
    st.stop()

# ---- 詩行表示 ----------------------------------------------------------------
st.subheader("詩行")
for i, line in enumerate(lines, 1):
    st.write(f"L{i:02d}: {line}")

poem_text = "\n".join(lines)

# ---- パラメータ --------------------------------------------------------------
st.markdown("### パラメータ")
colc1, colc2, colc3, colc4 = st.columns(4)
with colc1:
    use_raw_div = st.toggle("Divergenceを正規化しない（raw）", value=False)
with colc2:
    div_window = st.slider("Divergence window（過去行数）", 1, 7, 3, 1)
with colc3:
    decay_d = st.slider("Difficulty decay（余韻）", 0.0, 0.95, 0.50, 0.05)
with colc4:
    decay_r = st.slider("Resonance decay（余韻）", 0.0, 0.95, 0.50, 0.05)

# ---- GPT抽出 -----------------------------------------------------------------
difficult_prompt = (
    "以下の詩本文から『難解な語』を最大10語、抽出してください。"
    "哲学的・象徴的・抽象的な語を重視し、語のみを改行区切りで返してください。\n\n"
    f"{poem_text}"
)
resonance_prompt = (
    "以下の詩本文から『共鳴性の高い語』を最大10語、抽出してください。"
    "感覚的・感情的イメージを喚起する語を重視し、語のみを改行区切りで返してください。\n\n"
    f"{poem_text}"
)

col1, col2 = st.columns(2)
with col1:
    difficult_words = gpt_extract_words(difficult_prompt)
    st.subheader("難解語群（GPT抽出／本文依拠）")
    if difficult_words:
        st.table(pd.DataFrame(difficult_words, columns=["Word"]))
    else:
        st.info("難解語の自動抽出に失敗/空。空集合として解析を続行します。")
with col2:
    resonance_words = gpt_extract_words(resonance_prompt)
    st.subheader("共鳴語群（GPT抽出／本文依拠）")
    if resonance_words:
        st.table(pd.DataFrame(resonance_words, columns=["Word"]))
    else:
        st.info("共鳴語の自動抽出に失敗/空。空集合として解析を続行します。")

# ---- スコア算定 --------------------------------------------------------------
model = load_embedder()
divergence_scores = compute_divergence_scores(
    lines, model, window=div_window, normalize=(not use_raw_div)
)
difficulty_scores = compute_difficulty_scores(lines, difficult_words, decay=decay_d)
resonance_scores  = compute_resonance_scores_from_terms(lines, resonance_words, decay=decay_r)

df = pd.DataFrame({
    "Line": [f"L{i+1:02d}" for i in range(len(lines))],
    "Text": lines,
    "Divergence": divergence_scores,
    "Difficulty": difficulty_scores,
    "Resonance": resonance_scores
})

# ---- 表・DL ------------------------------------------------------------------
st.subheader("スコア一覧")
st.caption("Divergence は raw（1−cos）/ 正規化（0–1）をトグル。window=過去行平均、decay=余韻。")
st.dataframe(df, use_container_width=True)

csv = df.to_csv(index=False).encode("utf-8-sig")
st.download_button("CSVをダウンロード", csv, "scores.csv", "text/csv")

# ---- 波形 --------------------------------------------------------------------
st.subheader("波形可視化")
ylabel = "Divergence (raw 1−cos)" if use_raw_div else "Divergence (0–1 normalized)"
fig, ax = plt.subplots(figsize=(10, 5))
ax.plot(df["Line"], df["Divergence"], label="Divergence")
ax.plot(df["Line"], df["Difficulty"], label="Difficulty")
ax.plot(df["Line"], df["Resonance"], label="Resonance")
ax.set_xlabel("Line")
ax.set_ylabel(ylabel)
ax.legend()
ax.grid(True, linestyle="--", alpha=0.5)
st.pyplot(fig)

# ---- 四象限（raw）＋重心 -----------------------------------------------------
st.subheader("四象限マップ（raw）＋重心表示")
x = df["Difficulty"].values
y = df["Resonance"].values

# 点サイズ：Divergence（内部正規化してサイズ化、rawでも視認性のため）
dv = df["Divergence"].values
dv_norm = _safe_minmax_scale(dv) if dv.size else dv
sizes = 80 + 220 * dv_norm

cx, cy = float(np.mean(x)), float(np.mean(y))

fig2, ax2 = plt.subplots(figsize=(8, 6))
ax2.scatter(x, y, s=sizes, alpha=0.75, label="Lines")
ax2.scatter(cx, cy, c="red", s=200, marker="X", label="重心")

# 行ラベル（多い場合は重なるので軽量表示）
for xi, yi, label in zip(x, y, df["Line"].astype(str).values):
    ax2.annotate(label, (xi, yi), textcoords="offset points", xytext=(3, 3), fontsize=7)

ax2.set_xlabel("Difficulty (raw)")
ax2.set_ylabel("Resonance (raw)")
ax2.set_title("Quadrant Scatter (raw) with Centroid")
ax2.legend()
st.pyplot(fig2)

# ---- レポート出力 ------------------------------------------------------------
st.subheader("レポート出力")
# DOCX
try:
    docx_buf = build_docx_report(
        lines=lines,
        df=df,
        difficult_words=difficult_words,
        resonance_words=resonance_words,
        fig_wave=fig,       # 既存の波形図
        fig_quad=fig2,      # 既存の四象限図
        title="Poetic Waveform Report"
    )
    st.download_button(
        "DOCXレポートをダウンロード",
        data=docx_buf.getvalue(),
        file_name="poetic_waveform_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
except Exception as e:
    st.warning(f"DOCX生成に失敗しました: {e}")


# PDF
try:
    pdf_buf = build_pdf_report(
        lines=lines,
        df=df,
        fig_wave=fig,
        fig_quad=fig2,
        title="Poetic Waveform Report",
        font_family=used_family   # ←ここを追加
    )
    st.download_button(
        "PDFレポートをダウンロード",
        data=pdf_buf.getvalue(),
        file_name="poetic_waveform_report.pdf",
        mime="application/pdf",
    )
except Exception as e:
    st.warning(f"PDF生成に失敗しました: {e}")



